import os
from transformers import AutoTokenizer, AutoModel

# 获取文件路径
current_file_path = os.path.abspath(__file__)
project_root = os.path.dirname(os.path.dirname(current_file_path))
model_path = os.path.join(project_root, "models", "bge-small-zh-v1.5")

print(f"正在从本地加载模型，路径为: {model_path}")

# 配置文件里定义好的各级标题
from config_guide import (
    DOC_TITLE_BLACKBALL,
    CHAPTER_HEADERS,
    SUB_BLOCK_TRIGGERS,
    CATEGORY_MAP,
)


from docx import Document
from docx.oxml.ns import qn
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document as LangchainDocument
import re


DOCX_PATH = "../data/灵山胜境：历史、文化、景点特色与个性化游览指南.docx"
VECTOR_STORE_PATH = "../vector_store/lingshan"


# 1. 定义表格处理函数


def table_to_text(table) -> str:
    """把表格转成自然语言字符串，供合并节内容时使用"""
    rows = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])

    if not rows:
        return ""

    first_row = rows[0]
    is_header = len(rows) > 1 and all(len(cell) < 20 for cell in first_row)

    lines = []
    if is_header:
        headers = first_row
        for row in rows[1:]:
            pairs = []
            seen = set()
            for header, cell in zip(headers, row):
                if cell and cell not in seen:
                    pairs.append(f"{header}：{cell}")
                    seen.add(cell)
            if pairs:
                lines.append("；".join(pairs))
    else:
        for row in rows:
            seen = set()
            cells = [c for c in row if c and c not in seen and not seen.add(c)]
            if cells:
                lines.append("；".join(cells))

    return "\n".join(lines)


def table_to_rows(table) -> list[dict]:
    """
    把表格拆成行级别的列表，供构造子chunk使用。
    每个元素是 {"row_label": str, "row_text": str}
    """
    rows = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])

    if not rows:
        return []

    first_row = rows[0]
    is_header = len(rows) > 1 and all(len(cell) < 20 for cell in first_row)

    result = []
    if is_header:
        headers = first_row
        for row in rows[1:]:
            pairs = []
            seen = set()
            label = ""
            for i, (header, cell) in enumerate(zip(headers, row)):
                if cell and cell not in seen:
                    if i == 0:
                        label = cell
                    pairs.append(f"{header}：{cell}")
                    seen.add(cell)
            if pairs:
                result.append(
                    {
                        "row_label": label,
                        "row_text": "；".join(pairs),
                    }
                )
    else:
        for row in rows:
            seen = set()
            cells = [c for c in row if c and c not in seen and not seen.add(c)]
            if cells:
                result.append(
                    {
                        "row_label": cells[0],
                        "row_text": "；".join(cells),
                    }
                )

    return result


# 2. 提取原始block


def extract_raw_blocks(docx_path: str) -> list[dict]:
    """按文档原始顺序提取所有段落和表格"""
    doc = Document(docx_path)
    blocks = []

    # 建立表格元素的映射字典，为了保持原始顺序
    table_dict = {tbl._element: tbl for tbl in doc.tables}

    # 遍历文档的主体元素（包含段落和表格），保证提取顺序不乱
    for element in doc.element.body:
        if element.tag.endswith("}p"):
            # 如果是段落 (Paragraph)
            from docx.text.paragraph import Paragraph

            para = Paragraph(element, doc._part)
            text = para.text.strip()
            if text:
                blocks.append({"type": "text", "content": text})

        elif element.tag.endswith("}tbl"):
            # 如果是表格 (Table)
            if element in table_dict:
                table = table_dict[element]
                table_text = table_to_text(table)
                if table_text:
                    blocks.append({"type": "table", "content": table_text})

    return blocks


# 3. 标题判断、大类映射、合并为节


def is_heading(text: str) -> bool:
    """判断一段文本是否是标题"""
    if len(text) > 40:  # 太长的不是标题
        return False
    if text.endswith("：") or text.endswith(":"):  # 冒号结尾的不是标题
        return False
    if text.count("，") + text.count(",") > 1:  # 带很多逗号的不是标题
        return False
    content_patterns = (
        [  # 再手动排除一些不是标题的（可能不太灵活，我这里为了方便直接写死了）
            r"^\d{4}",
            r"^在",
            r"^路线规划",
            r"^讲解重点",
            r"^特色体验",
            r"^[•\-\*]",
            r"^参与",
            r"^观赏",
            r"^漫步",
        ]
    )
    for pattern in content_patterns:
        if re.match(pattern, text):
            return False
    if "元/" in text or "元起" in text:
        return False
    return True


def get_category(section_title: str) -> str:
    """根据节标题返回大类名称"""
    for keyword, category in CATEGORY_MAP.items():
        if keyword in section_title:
            return category
    return "其他"


def merge_blocks_into_sections(raw_blocks: list[dict]) -> list[dict]:
    """
    实现精确的层级合并：
    1. 过滤掉文档总标题。
    2. 识别“章节”级标题，并将其作为后续小节的父级面包屑。
    """
    sections = []

    current_macro_title = ""  # 记录当前的章节父标题
    current_title = "概述"
    current_lines = []
    current_text_lines = []
    current_has_table = False
    current_types = set()

    tbl_counter = 0

    current_table_indices = []

    def flush():
        nonlocal current_macro_title
        if not current_lines:
            return

        # A. 过滤逻辑：如果是文档总标题，直接丢弃，且不更新任何状态
        if DOC_TITLE_BLACKBALL in current_title:
            return

        # B. 章节识别逻辑
        # 检查当前标题是否属于预设的五个大章节之一
        is_chapter = any(ch in current_title for ch in CHAPTER_HEADERS)

        if is_chapter:
            # 如果是大章节，它本身就是最高级
            current_macro_title = current_title
            breadcrumb = current_title
        elif current_macro_title:
            # 如果是普通小节，则拼接之前记住的大章节标题
            breadcrumb = f"{current_macro_title} > {current_title}"
        else:
            breadcrumb = current_title

        # C. 过滤空壳标题
        # 如果只有一行标题（没有正文也没表格），则只更新 macro_title，不生成独立文档
        if len(current_lines) == 1 and not current_has_table:
            return

        # D. 写入面包屑并封装
        current_lines[0] = breadcrumb

        if current_text_lines:
            current_text_lines[0] = breadcrumb
        sections.append(
            {
                "section_title": breadcrumb,
                "content": "\n".join(current_lines),
                "text_content": "\n".join(current_text_lines),
                "has_table": current_has_table,
                "source_types": list(current_types),
                "table_indices": list(current_table_indices),
            }
        )

    for block in raw_blocks:
        text = block["content"]
        if block["type"] == "text" and is_heading(text):
            flush()
            current_title = text
            current_lines = [text]
            current_text_lines = [text]
            current_has_table = False
            current_types = {"text"}
            current_table_indices = []
        else:
            current_lines.append(text)
            current_types.add(block["type"])
            if block["type"] == "text":
                current_text_lines.append(text)
            elif block["type"] == "table":
                current_has_table = True
                current_table_indices.append(tbl_counter)
                tbl_counter += 1

    flush()
    return sections


# 4. 构造父子文档


def sections_to_documents_hierarchical(
    sections: list[dict],
    source_filename: str,
    raw_blocks: list[dict],
    doc_tables: list,
) -> tuple[list[LangchainDocument], list[LangchainDocument]]:
    """
    构造两级文档：
    - 父文档：完整节内容，存入 guide_parents，送给LLM使用
    - 子文档：细粒度chunk，存入 guide_children，用于相似度检索
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=200,
        chunk_overlap=30,
        separators=["\n", "。", "；", "，", " ", ""],
    )

    parent_docs = []
    child_docs = []

    for i, section in enumerate(sections):
        parent_id = f"{source_filename}::section::{i}"
        category = get_category(section["section_title"])
        breadcrumb = section["section_title"]

        # A. 父文档：完整内容，不切分
        parent_docs.append(
            LangchainDocument(
                page_content=section["content"],
                metadata={
                    "source": source_filename,
                    "type": "guide",
                    "section": section["section_title"],
                    "category": category,
                    "parent_id": parent_id,
                    "is_parent": "true",
                },
            )
        )

        # B. 构造文本子文档（基于业务标识词聚合）
        content_lines = section["content"].split("\n")

        body_lines = content_lines[1:] if len(content_lines) > 1 else []

        grouped_chunks = []
        current_sub_title = ""
        current_buffer = []

        def flush_buffer():
            if current_buffer:
                grouped_chunks.append(
                    {
                        "sub_title": current_sub_title,
                        "content": "\n".join(current_buffer),
                    }
                )
                current_buffer.clear()

        for line in body_lines:
            line = line.strip()
            if not line:
                continue

            # 检查当前行是否触发了新的子模块
            is_trigger = False
            for trigger in SUB_BLOCK_TRIGGERS:
                # 兼容中文冒号、英文冒号以及空格
                if line.startswith(trigger) and (
                    len(line) == len(trigger)
                    or line[len(trigger) :].strip().startswith("：")
                    or line[len(trigger) :].strip().startswith(":")
                ):
                    flush_buffer()
                    current_sub_title = trigger
                    current_buffer.append(line)  # 将标题行也保留在内容中
                    is_trigger = True
                    break

            if not is_trigger:
                # 如果不是新标题，就继续往当前的 buffer 里塞
                current_buffer.append(line)

        # 遍历结束后，把最后一个 buffer 也推入
        flush_buffer()

        # 生成最终的 LangchainDocument
        for j, chunk_data in enumerate(grouped_chunks):
            sub_title = chunk_data["sub_title"]
            chunk_content = chunk_data["content"]

            # 将子标题加入面包屑中
            if sub_title:
                child_breadcrumb = f"{breadcrumb} > {sub_title}"
            else:
                child_breadcrumb = breadcrumb

            # 最终的检索文本结构：
            final_page_content = f"{child_breadcrumb}\n{chunk_content}"

            child_docs.append(
                LangchainDocument(
                    page_content=final_page_content,
                    metadata={
                        "source": source_filename,
                        "type": "guide",
                        "section": breadcrumb,
                        "category": category,
                        "parent_id": parent_id,
                        "is_parent": "false",
                        "content_type": "text_segment",
                        "chunk_index": j,
                    },
                )
            )
        # C. 构造“表格”子文档（实现“一行一拆”）
        if section["has_table"]:
            # 找出属于这个 section 的所有 table 对象
            section_tables = [
                doc_tables[idx]
                for idx in section["table_indices"]
                if idx < len(doc_tables)
            ]

            for tbl in section_tables:
                # 获取表格的每一行
                rows_data = table_to_rows(tbl)
                for k, row in enumerate(rows_data):
                    # 强制每一行表格都带上面包屑标题
                    # 这里直接使用上面已经定义好的 breadcrumb 变量
                    table_row_content = f"{breadcrumb} —— {row['row_text']}"

                    child_docs.append(
                        LangchainDocument(
                            page_content=table_row_content,
                            metadata={
                                "parent_id": parent_id,
                                "is_parent": "false",
                                "content_type": "table_row",
                                "section": breadcrumb,
                                "category": category,
                                "source": source_filename,
                                "row_index": k,
                            },
                        )
                    )

    return parent_docs, child_docs


# 5. 写入Chroma


def ingest(docx_path: str, vector_store_path: str):
    print("第1步：提取原始block...")
    raw_blocks = extract_raw_blocks(docx_path)
    print(f"  → 共 {len(raw_blocks)} 个block")

    print("第2步：合并为语义节...")
    sections = merge_blocks_into_sections(raw_blocks)
    print(f"  → 共 {len(sections)} 个节")

    print("第3步：构造父子文档...")
    doc = Document(docx_path)
    parent_docs, child_docs = sections_to_documents_hierarchical(
        sections=sections,
        source_filename=os.path.basename(docx_path),
        raw_blocks=raw_blocks,
        doc_tables=doc.tables,
    )
    print(f"  → 父文档 {len(parent_docs)} 个，子文档 {len(child_docs)} 个")

    print("第4步：加载embedding模型，路径为：{model_path}...")
    embeddings = HuggingFaceEmbeddings(
        model_name=model_path,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )

    print("第5步：写入向量数据库...")
    os.makedirs(vector_store_path, exist_ok=True)

    Chroma.from_documents(
        documents=child_docs,
        embedding=embeddings,
        persist_directory=vector_store_path,
        collection_name="guide_children",
    )
    print("  → guide_children 写入完成")

    Chroma.from_documents(
        documents=parent_docs,
        embedding=embeddings,
        persist_directory=vector_store_path,
        collection_name="guide_parents",
    )
    print("  → guide_parents 写入完成")

    print("全部完成！")


# 6. 调试和验证


def preview_sections(docx_path: str):
    """预览合并后的节，确认标题和内容是否正确"""
    raw_blocks = extract_raw_blocks(docx_path)
    sections = merge_blocks_into_sections(raw_blocks)
    for i, sec in enumerate(sections):
        types_label = "+".join(sec["source_types"])
        print(f"[节 {i}] [{types_label}] 标题：{sec['section_title']}")
        print(sec["content"])
        print("---")


def verify_vectorstore(vector_store_path: str):
    """写入完成后验证父子检索是否正常工作"""
    print("加载embedding模型...")
    embeddings = HuggingFaceEmbeddings(
        model_name=model_path,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )

    child_store = Chroma(
        persist_directory=vector_store_path,
        embedding_function=embeddings,
        collection_name="guide_children",
    )
    parent_store = Chroma(
        persist_directory=vector_store_path,
        embedding_function=embeddings,
        collection_name="guide_parents",
    )

    query = "什么是五印坛城"
    print(f"\n测试问题：{query}")
    print("=" * 40)

    child_results = child_store.similarity_search(query, k=3)
    print("命中的子chunk：")
    for doc in child_results:
        print(f"  所属节：{doc.metadata['section']}")
        print(f"  内容：{doc.page_content}")
        print(f"  parent_id：{doc.metadata['parent_id']}")
        print()

    print("对应的父节完整内容：")
    seen = set()
    for doc in child_results:
        pid = doc.metadata["parent_id"]
        if pid in seen:
            continue
        seen.add(pid)
        result = parent_store.get(where={"parent_id": pid})
        if result["documents"]:
            print(f"  父节标题：{result['metadatas'][0]['section']}")
            print(f"  父节内容：{result['documents']}")
            print()


if __name__ == "__main__":
    # 按顺序执行，每步确认没问题再进行下一步

    # 第一步：预览节的合并效果
    # preview_sections(DOCX_PATH)

    # 第二步：写入向量数据库
    ingest(DOCX_PATH, VECTOR_STORE_PATH)

    # 第三步：验证父子检索
    # verify_vectorstore(VECTOR_STORE_PATH)
