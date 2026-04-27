from docx import Document
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document as LangchainDocument
import os

from config_dataset import TABLE_NAMES, SEARCHABLE_COLUMNS

# 复用 ingest_guide.py 里的这两个函数
from ingest_guide import table_to_text, table_to_rows

DOCX_PATH = "data/灵山胜境 景点结构化数据集.docx"
VECTOR_STORE_PATH = "vector_store/lingshan"  # 和指南用同一个路径


def process_dataset(docx_path: str) -> tuple[list, list]:
    doc = Document(docx_path)
    parent_docs = []
    child_docs = []

    # 数据集里有两张表：灵山胜境和拈花湾
    # 分别处理，用 source_table 区分
    # 开头的那段话没放进来处理
    

    for tbl_idx, table in enumerate(doc.tables):
        if tbl_idx >= len(TABLE_NAMES):
            break

        table_name = TABLE_NAMES[tbl_idx]
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])

        if not rows:
            continue

        headers = rows[0]   # 第一行是表头

        for row_idx, row in enumerate(rows[1:], start=1):
            # 跳过数据说明行（第一行数据是说明，不是景点）
            if not any(row):
                continue

            # 找到景点名称列（第3列，index=2）
            spot_name = row[2] if len(row) > 2 else f"景点{row_idx}"
            if not spot_name or spot_name == headers[2]:
                continue

            parent_id = f"dataset::{table_name}::{spot_name}"

            # A. 父文档：完整一行，所有列拼在一起
            full_text_parts = []
            for header, cell in zip(headers, row):
                if cell and cell not in headers:  # 跳过表头重复内容
                    full_text_parts.append(f"{header}：{cell}")

            parent_content = f"{spot_name}\n" + "\n".join(full_text_parts)

            parent_docs.append(
                LangchainDocument(
                    page_content=parent_content,
                    metadata={
                        "source": os.path.basename(docx_path),
                        "type": "dataset",          # 区别于 guide
                        "source_table": table_name,
                        "spot_name": spot_name,
                        "parent_id": parent_id,
                        "is_parent": "true",
                        "category": "景点介绍",
                    }
                )
            )

            # B. 子文档：每个可检索列单独一个chunk
            for header, cell in zip(headers, row):
                if header not in SEARCHABLE_COLUMNS:
                    continue
                if not cell or len(cell) < 10:  # 太短的跳过
                    continue

                child_content = f"{spot_name} —— {header}：{cell}"

                child_docs.append(
                    LangchainDocument(
                        page_content=child_content,
                        metadata={
                            "source": os.path.basename(docx_path),
                            "type": "dataset",
                            "source_table": table_name,
                            "spot_name": spot_name,
                            "column": header,       # 记录来自哪一列
                            "parent_id": parent_id,
                            "is_parent": "false",
                            "category": "景点介绍",
                        }
                    )
                )

    return parent_docs, child_docs


def ingest_dataset(docx_path: str, vector_store_path: str):
    print("处理数据集文档...")
    parent_docs, child_docs = process_dataset(docx_path)
    print(f"  → 父文档 {len(parent_docs)} 个，子文档 {len(child_docs)} 个")

    model_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "models", "bge-small-zh-v1.5"
    )

    print("加载embedding模型...")
    embeddings = HuggingFaceEmbeddings(
        model_name=model_path,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )

    print("写入向量数据库...")
    # 这里用 add_documents 而不是 from_documents
    # 因为 collection 已经被 ingest_guide.py 创建过了（顺序上先运行的它），不能重复创建

    child_store = Chroma(
        persist_directory=vector_store_path,
        embedding_function=embeddings,
        collection_name="guide_children",
    )
    child_store.add_documents(child_docs)
    print("  → guide_children 追加完成")

    parent_store = Chroma(
        persist_directory=vector_store_path,
        embedding_function=embeddings,
        collection_name="guide_parents",
    )
    parent_store.add_documents(parent_docs)
    print("  → guide_parents 追加完成")

    print("完成！")


if __name__ == "__main__":
    ingest_dataset(DOCX_PATH, VECTOR_STORE_PATH)